"""
SLS Studio Legale - Document Formatter v2
Formatta atti legali secondo il modello dello studio.
Integra Claude AI per interpretazione intelligente del contenuto.
"""

import os
import re
import uuid
import json
from datetime import datetime
from flask import Flask, render_template, request, send_file, flash, redirect, url_for, session
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import httpx

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'sls-studio-legale-secret-key-change-in-production')
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['OUTPUT_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'outputs')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'docx'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_docx(filepath):
    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            fmt = []
            for run in para.runs:
                if run.bold: fmt.append('bold')
                if run.italic: fmt.append('italic')
                if run.underline: fmt.append('underline')
            paragraphs.append({
                'text': para.text,
                'original_format': ','.join(set(fmt)) if fmt else 'normal',
                'style': para.style.name if para.style else 'Normal'
            })
    return paragraphs


def call_claude_ai(text_content):
    if not ANTHROPIC_API_KEY:
        return None

    paragraphs_text = ""
    for i, p in enumerate(text_content):
        paragraphs_text += f"[{i}] {p['text'][:200]}\n"

    prompt = f"""Sei un assistente legale esperto nella formattazione di atti giudiziari italiani.

Analizza i seguenti paragrafi di un atto legale e classifica OGNUNO con il tipo di formattazione appropriato.

I tipi possibili sono:
- "intestazione_studio" = iniziali o nome dello studio/avvocato in alto
- "autorita" = nome del tribunale/corte
- "tipo_atto" = tipo di documento (MEMORIA DIFENSIVA, ATTO DI CITAZIONE, RICORSO, ecc.)
- "riferimento_rg" = riferimento al numero R.G.
- "separatore" = linea di separazione (*** o simili)
- "nell_interesse" = "Nell'interesse di:" o simili
- "nome_parte" = nome e dati anagrafici della parte assistita
- "qualifica_parte" = ruolo processuale (es. "- convenuto -")
- "contro" = "contro:" o "nei confronti di:"
- "nome_controparte" = nome e dati della controparte
- "qualifica_controparte" = ruolo processuale controparte
- "titolo_sezione" = titoli come PREMESSO CHE, MOTIVI, CONCLUSIONI, IN FATTO, IN DIRITTO
- "titolo_numerato" = titoli numerati (I - ..., II - ..., 1) ...)
- "corpo" = testo normale del corpo
- "corpo_rientrato" = sotto-argomenti che necessitano rientro prima riga
- "sotto_conclusione" = In via pregiudiziale, In via principale, In subordine
- "si_producono" = "Si producono:" o "Si allegano:"
- "documento_lista" = singolo documento nella lista produzioni
- "luogo_data" = luogo e data
- "firma" = firma dell'avvocato
- "valore_causa" = dichiarazione contributo unificato

Rispondi SOLO con un JSON array: [{{"index": 0, "type": "..."}}, ...]

PARAGRAFI:
{paragraphs_text}"""

    try:
        response = httpx.post(
            'https://api.anthropic.com/v1/messages',
            headers={
                'x-api-key': ANTHROPIC_API_KEY,
                'anthropic-version': '2023-06-01',
                'content-type': 'application/json',
            },
            json={
                'model': 'claude-sonnet-4-20250514',
                'max_tokens': 4096,
                'messages': [{'role': 'user', 'content': prompt}]
            },
            timeout=60.0
        )
        result = response.json()
        text = result['content'][0]['text']
        json_match = re.search(r'\[.*\]', text, re.DOTALL)
        if json_match:
            classifications = json.loads(json_match.group())
            return {item['index']: item['type'] for item in classifications}
    except Exception as e:
        print(f"Errore Claude AI: {e}")
    return None


def classify_paragraph_rules(text, index, total):
    t = text.strip()
    t_upper = t.upper()
    t_lower = t.lower()
    is_start = index < 5
    is_end = index > total - 5

    if t in ('*** *** ***', '***', '* * *', '***   ***   ***'):
        return 'separatore'
    if is_start and len(t) <= 5 and t.replace('.', '').replace(' ', '').isalpha():
        return 'intestazione_studio'
    if 'TRIBUNALE' in t_upper or 'CORTE' in t_upper or 'GIUDICE DI PACE' in t_upper or 'CGT' in t_upper:
        if len(t) < 80:
            return 'autorita'
    if any(x in t_upper for x in ['MEMORIA', 'ATTO DI CITAZIONE', 'RICORSO', 'COMPARSA', 'APPELLO', 'RECLAMO', 'OPPOSIZIONE']):
        if len(t) < 100:
            return 'tipo_atto'
    if ('R.G.' in t or 'r.g.' in t_lower) and len(t) < 100:
        return 'riferimento_rg'
    if t_lower.startswith("nell'interesse") or t_lower.startswith("nell'interesse"):
        return 'nell_interesse'
    if t_lower.strip().startswith('contro') and len(t) < 30:
        return 'contro'
    if t_lower.strip().startswith('nei confronti di'):
        return 'contro'
    if re.search(r'-\s*(convenut|attor|ricorrent|resistente|oppost|appellante|appellat)', t_lower):
        return 'qualifica_parte' if index < total / 2 else 'qualifica_controparte'
    if any(x in t_upper for x in ['PREMESSO CHE', 'IN FATTO', 'IN DIRITTO', 'MOTIVI', 'CONCLUSIONI', 'P.Q.M.', 'CHIEDE']) and len(t) < 50:
        return 'titolo_sezione'
    if re.match(r'^[IVX]+[\s\u2013\-]', t) or re.match(r'^[A-Z]\)', t):
        return 'titolo_numerato'
    if any(x in t_lower for x in ['in via pregiudiziale', 'in via principale', 'in subordine', 'in via istruttoria', 'in via gradata']) and len(t) < 80:
        return 'sotto_conclusione'
    if t_lower.startswith('si produc') or t_lower.startswith('si alleg'):
        return 'si_producono'
    if 'contributo unificato' in t_lower:
        return 'valore_causa'
    if is_end and ('avv.' in t_lower or 'dott.' in t_lower) and len(t) < 60:
        return 'firma'
    if is_end and ',' in t and any(c.isdigit() for c in t) and len(t) < 40:
        return 'luogo_data'
    if 'C.F.' in t or 'P.IVA' in t_upper:
        return 'nome_parte'
    return 'corpo'


def format_document(paragraphs, classifications):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(4.1)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(3.26)

    FORMAT_MAP = {
        'intestazione_studio': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY},
        'autorita': {'align': WD_ALIGN_PARAGRAPH.CENTER, 'bold': True},
        'tipo_atto': {'align': WD_ALIGN_PARAGRAPH.CENTER, 'bold': True, 'underline': True},
        'riferimento_rg': {'align': WD_ALIGN_PARAGRAPH.CENTER, 'bold': True},
        'separatore': {'align': WD_ALIGN_PARAGRAPH.CENTER, 'text_override': '*** *** ***'},
        'nell_interesse': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY, 'underline': True},
        'nome_parte': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY, 'bold': True},
        'nome_controparte': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY, 'bold': True},
        'qualifica_parte': {'align': WD_ALIGN_PARAGRAPH.RIGHT, 'bold': True, 'italic': True},
        'qualifica_controparte': {'align': WD_ALIGN_PARAGRAPH.RIGHT, 'bold': True, 'italic': True},
        'contro': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY, 'underline': True},
        'titolo_sezione': {'align': WD_ALIGN_PARAGRAPH.CENTER, 'bold': True, 'underline': True},
        'titolo_numerato': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY, 'bold': True},
        'corpo': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY},
        'corpo_rientrato': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY, 'indent': Cm(1.05)},
        'sotto_conclusione': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY, 'bold': True},
        'si_producono': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY, 'underline': True},
        'documento_lista': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY},
        'valore_causa': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY, 'italic': True},
        'luogo_data': {'align': WD_ALIGN_PARAGRAPH.JUSTIFY},
        'firma': {'align': WD_ALIGN_PARAGRAPH.RIGHT},
    }

    for i, para_data in enumerate(paragraphs):
        text = para_data['text']
        para_type = classifications.get(i, 'corpo')
        fmt = FORMAT_MAP.get(para_type, FORMAT_MAP['corpo'])

        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.line_spacing = 1.5
        pf.alignment = fmt.get('align', WD_ALIGN_PARAGRAPH.JUSTIFY)

        if 'indent' in fmt:
            pf.first_line_indent = fmt['indent']

        display_text = fmt.get('text_override', text)
        run = p.add_run(display_text)
        run.font.size = Pt(12)

        if fmt.get('bold'):
            run.bold = True
        if fmt.get('italic'):
            run.italic = True
        if fmt.get('underline'):
            run.underline = True

        # Per il corpo, preserva formattazione originale
        if para_type == 'corpo' and para_data.get('original_format', 'normal') != 'normal':
            ofmt = para_data['original_format']
            if 'bold' in ofmt: run.bold = True
            if 'italic' in ofmt: run.italic = True
            if 'underline' in ofmt: run.underline = True

    return doc


@app.route('/')
def index():
    ai_available = bool(ANTHROPIC_API_KEY)
    return render_template('index.html', step=1, ai_available=ai_available)


@app.route('/format', methods=['POST'])
def format_doc():
    if 'document' not in request.files:
        flash('Nessun file selezionato!', 'error')
        return redirect(url_for('index'))

    file = request.files['document']
    if file.filename == '' or not allowed_file(file.filename):
        flash('Solo file .docx sono accettati!', 'error')
        return redirect(url_for('index'))

    doc_name = f"doc_{uuid.uuid4().hex}_{secure_filename(file.filename)}"
    doc_path = os.path.join(app.config['UPLOAD_FOLDER'], doc_name)
    file.save(doc_path)

    try:
        paragraphs = extract_text_from_docx(doc_path)
        if not paragraphs:
            flash('Il documento sembra vuoto!', 'error')
            return redirect(url_for('index'))

        ai_classifications = None
        used_ai = False

        if ANTHROPIC_API_KEY:
            ai_classifications = call_claude_ai(paragraphs)
            if ai_classifications:
                used_ai = True

        if not ai_classifications:
            ai_classifications = {}
            for idx, p in enumerate(paragraphs):
                ai_classifications[idx] = classify_paragraph_rules(p['text'], idx, len(paragraphs))

        output_doc = format_document(paragraphs, ai_classifications)

        output_name = f"FORMATTATO_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{secure_filename(file.filename)}"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_name)
        output_doc.save(output_path)

        session['output_file'] = output_name
        session['output_original'] = file.filename

        report_lines = []
        for idx, p in enumerate(paragraphs):
            ptype = ai_classifications.get(idx, 'corpo')
            preview = p['text'][:60] + '...' if len(p['text']) > 60 else p['text']
            report_lines.append(f"[{ptype:>25}]  {preview}")

        return render_template('index.html', step=2,
                               output_name=file.filename,
                               report='\n'.join(report_lines),
                               used_ai=used_ai,
                               total=len(paragraphs))

    except Exception as e:
        flash(f'Errore: {str(e)}', 'error')
        return redirect(url_for('index'))


@app.route('/download')
def download():
    if 'output_file' not in session:
        flash('Nessun file da scaricare.', 'error')
        return redirect(url_for('index'))
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], session['output_file'])
    return send_file(output_path, as_attachment=True,
                     download_name=f"FORMATTATO_{session.get('output_original', 'documento.docx')}")


@app.route('/reset')
def reset():
    session.clear()
    return redirect(url_for('index'))


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=True, host='0.0.0.0', port=port)
